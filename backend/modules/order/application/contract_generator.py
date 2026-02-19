"""
Contract Generator Service
Generates .docx contract from Word template 'HDDV Giao Tuyet Template.docx'
by filling in order data into placeholder paragraphs.
Preserves all formatting, fonts, and styles from the template.
"""

import copy
import io
import re
from datetime import datetime, timedelta
from decimal import Decimal
from pathlib import Path
from typing import Optional

from docx import Document

# Path to the contract template (inside backend/templates/)
TEMPLATE_PATH = Path(__file__).resolve().parents[3] / "templates" / "contract_template.docx"

# Word XML namespace
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ─── Vietnamese number-to-words helper ───────────────────────────

_ONES = [
    '', 'một', 'hai', 'ba', 'bốn', 'năm', 'sáu', 'bảy', 'tám', 'chín'
]
_TENS_SPECIAL = {
    0: 'mười',
    1: 'mười một',
    5: 'mười lăm',
}
_GROUPS = ['', 'nghìn', 'triệu', 'tỷ']


def _read_three_digits(n: int) -> str:
    """Convert a 3-digit number to Vietnamese words."""
    if n == 0:
        return ''

    hundreds = n // 100
    tens = (n % 100) // 10
    ones = n % 10

    parts = []

    if hundreds > 0:
        parts.append(f'{_ONES[hundreds]} trăm')
        if tens == 0 and ones > 0:
            parts.append('lẻ')
    
    if tens > 0:
        if tens == 1:
            parts.append('mười')
        else:
            parts.append(f'{_ONES[tens]} mươi')
        
        if ones == 1 and tens > 1:
            parts.append('mốt')
        elif ones == 5 and tens > 0:
            parts.append('lăm')
        elif ones > 0:
            parts.append(_ONES[ones])
    elif ones > 0 and hundreds > 0:
        parts.append(_ONES[ones])
    elif ones > 0 and hundreds == 0:
        parts.append(_ONES[ones])

    return ' '.join(parts)


def number_to_vietnamese_words(amount: int) -> str:
    """Convert a number to Vietnamese words (for currency).
    
    Example: 15_501_400 -> 'Mười lăm triệu năm trăm lẻ một nghìn bốn trăm đồng'
    """
    if amount == 0:
        return 'Không đồng'

    # Split into groups of 3 from right
    groups = []
    n = abs(amount)
    while n > 0:
        groups.append(n % 1000)
        n //= 1000

    parts = []
    for i in range(len(groups) - 1, -1, -1):
        text = _read_three_digits(groups[i])
        if text:
            if _GROUPS[i]:
                parts.append(f'{text} {_GROUPS[i]}')
            else:
                parts.append(text)

    result = ' '.join(parts) + ' đồng'
    # Capitalize first letter
    return result[0].upper() + result[1:]


def format_vnd(amount) -> str:
    """Format amount as Vietnamese currency string.
    
    Example: 15501400 -> '15.501.400'
    """
    if amount is None:
        return '0'
    n = int(amount)
    # Use locale-style formatting with dots
    formatted = f'{n:,}'.replace(',', '.')
    return formatted


# ─── Paragraph text replacement helpers ──────────────────────────

def _replace_paragraph_text(paragraph, new_full_text: str) -> None:
    """Replace all text in a paragraph while keeping first run's formatting."""
    runs = paragraph.runs
    if not runs:
        return
    
    # Keep first run, set its text to the new text
    runs[0].text = new_full_text
    
    # Remove all other runs
    for run in runs[1:]:
        run._element.getparent().remove(run._element)


def _append_to_paragraph(paragraph, value: str) -> None:
    """Append value text after the existing label text in a paragraph.
    
    For paragraphs like 'Ông/Bà: [tab]' → 'Ông/Bà: Nguyễn Văn A'
    """
    runs = paragraph.runs
    if not runs:
        return
    
    # Reconstruct: keep the label part, append value
    full_text = paragraph.text
    
    # Find the colon or label ending position
    # Then rebuild: label + value
    # Strategy: clear all runs after the label, set last meaningful run's text
    
    # Find the label text (everything up to and including ':' + space)
    colon_idx = full_text.find(':')
    if colon_idx >= 0:
        label = full_text[:colon_idx + 1] + ' '
        new_text = label + value
    else:
        new_text = full_text + ' ' + value
    
    _replace_paragraph_text(paragraph, new_text)


# ─── Main generator function ────────────────────────────────────

class ContractData:
    """Data needed to fill the contract template."""
    def __init__(
        self,
        order_code: str,
        customer_name: str,
        customer_phone: str,
        customer_address: str,
        event_date: Optional[datetime] = None,
        event_time: Optional[str] = None,
        event_address: str = '',
        table_count: int = 0,
        dish_names: list = None,
        unit_price_per_table: int = 0,
        total_amount: int = 0,
        deposit_amount: int = 0,
        remaining_amount: int = 0,
        services_included: dict = None,
    ):
        self.order_code = order_code
        self.customer_name = customer_name or ''
        self.customer_phone = customer_phone or ''
        self.customer_address = customer_address or ''
        self.event_date = event_date
        self.event_time = event_time or ''
        self.event_address = event_address or ''
        self.table_count = table_count
        self.dish_names = dish_names or []
        self.unit_price_per_table = unit_price_per_table
        self.total_amount = total_amount
        self.deposit_amount = deposit_amount
        self.remaining_amount = remaining_amount
        self.services_included = services_included or {}


def generate_contract_docx(data: ContractData) -> io.BytesIO:
    """
    Generate a contract .docx file by filling in the template with order data.

    Args:
        data: ContractData with all order information

    Returns:
        BytesIO buffer containing the generated .docx file
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Contract template not found at: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))
    paragraphs = doc.paragraphs

    # ── 1. Contract Number (Table 0, Cell 0, Paragraph 2: "Số :") ──
    try:
        table0 = doc.tables[0]
        cell_0_0 = table0.rows[0].cells[0]
        for p in cell_0_0.paragraphs:
            if 'Số' in p.text and ':' in p.text:
                _replace_paragraph_text(p, f'Số : {data.order_code}')
                break
    except (IndexError, AttributeError):
        pass

    # ── 2. Party date line (P6: "Hôm nay, vào, tại:  chúng tôi gồm:") ──
    for p in paragraphs:
        if 'Hôm nay' in p.text and 'chúng tôi gồm' in p.text:
            # BUGFIX: "Hôm nay" = ngày in hợp đồng (today), KHÔNG phải ngày đãi tiệc
            today_str = datetime.now().strftime('%d/%m/%Y')
            _replace_paragraph_text(p, f'Hôm nay, vào ngày {today_str}, tại TP Hồ Chí Minh, chúng tôi gồm:')
            break

    # ── 3. Customer Info (P9-P12) ──
    for p in paragraphs:
        text = p.text.strip()
        
        if text.startswith('Ông/Bà:'):
            _replace_paragraph_text(p, f'Ông/Bà: {data.customer_name}')
        
        elif text.startswith('Địa chỉ:') and 'trụ sở' not in text and 'đãi tiệc' not in text:
            # Customer address (P10), not Party B address or event address
            _replace_paragraph_text(p, f'Địa chỉ: {data.customer_address}')
        
        elif text.startswith('Điện thoại:') and '0902786689' not in text:
            # Customer phone (P12), not Party B phone
            _replace_paragraph_text(p, f'Điện thoại: {data.customer_phone}')

    # ── 4. Event Details (P21-P24) ──
    for p in paragraphs:
        text = p.text.strip()
        
        # Party time
        if 'Thời gian đãi tiệc' in text and 'kết thúc' not in text:
            if data.event_date and data.event_time:
                date_str = data.event_date.strftime('%d/%m/%Y')
                _replace_paragraph_text(p, f'Thời gian đãi tiệc : {data.event_time} ngày {date_str}')
            elif data.event_time:
                _replace_paragraph_text(p, f'Thời gian đãi tiệc : {data.event_time}')
        
        # End time = party time + 4 hours
        elif 'Thời gian kết thúc tiệc' in text:
            if data.event_date and data.event_time:
                try:
                    h, m = map(int, data.event_time.split(':'))
                    end_time = f'{(h + 4) % 24}:{m:02d}'
                    date_str = data.event_date.strftime('%d/%m/%Y')
                    _replace_paragraph_text(p, f'Thời gian kết thúc tiệc: {end_time} ngày {date_str}')
                except (ValueError, AttributeError):
                    pass
        
        # Event address
        elif 'Địa chỉ đãi tiệc' in text:
            _replace_paragraph_text(p, f'Địa chỉ đãi tiệc: {data.event_address}')
        
        # Table count
        elif 'Số Bàn Tiệc' in text:
            _replace_paragraph_text(p, f'Số Bàn Tiệc : {data.table_count} bàn')

    # ── 5. Menu Items (P26-P31: placeholder tab lines after "Thực Đơn :") ──
    # Find the "Thực Đơn" paragraph and the dish placeholder paragraphs
    menu_start_idx = None
    for i, p in enumerate(paragraphs):
        if 'Thực Đơn' in p.text and 'bao gồm' in p.text:
            menu_start_idx = i
            break

    if menu_start_idx is not None:
        # Template has 6 placeholder paragraphs (P26-P31) after "Thực Đơn" (P25)
        # These are empty/tab-only paragraphs with "List Paragraph" style
        placeholder_start = menu_start_idx + 1
        placeholder_end = placeholder_start + 6  # 6 slots

        # Find "Đơn Giá" paragraph to know where placeholders end
        don_gia_idx = None
        for i, p in enumerate(paragraphs):
            if 'Đơn Giá' in p.text:
                don_gia_idx = i
                break

        if don_gia_idx:
            actual_placeholder_count = don_gia_idx - placeholder_start
        else:
            actual_placeholder_count = 6

        # Fill in dish names, then REMOVE unused placeholder paragraphs
        # BUGFIX: Must remove from XML, not just clear text — Word auto-numbering
        # still shows empty numbered items (4. 5. 6. 7.) if paragraphs exist
        paragraphs_to_remove = []
        for i in range(actual_placeholder_count):
            p_idx = placeholder_start + i
            if p_idx < len(paragraphs):
                if i < len(data.dish_names):
                    _replace_paragraph_text(paragraphs[p_idx], data.dish_names[i])
                else:
                    # Mark for removal (can't modify list while iterating)
                    paragraphs_to_remove.append(paragraphs[p_idx])

        # Remove unused placeholder paragraphs from XML tree
        for p in paragraphs_to_remove:
            p._element.getparent().remove(p._element)

    # ── 6. Pricing (P32-P34) ──
    for p in paragraphs:
        text = p.text.strip()
        
        # Unit price per table
        if 'Đơn Giá' in text:
            _replace_paragraph_text(p, f'Đơn Giá : {format_vnd(data.unit_price_per_table)} đ')
        
        # Total amount
        elif 'Tổng Cộng Thanh Toán' in text:
            _replace_paragraph_text(p, f'Tổng Cộng Thanh Toán : {format_vnd(data.total_amount)} đ')
        
        # Total in words
        elif 'Viết Bằng chữ' in text:
            _replace_paragraph_text(p, f'Viết Bằng chữ : {number_to_vietnamese_words(data.total_amount)}')

    # ── 7. Payment Terms (P48-P51) ──
    payment_1_done = False
    payment_1_words_done = False
    for p in paragraphs:
        text = p.text.strip()
        
        # Payment 1 (deposit)
        if 'Lần 1' in text and not payment_1_done:
            _replace_paragraph_text(
                p,
                f'Lần 1 : {format_vnd(data.deposit_amount)} đ khi ký hợp đồng.'
            )
            payment_1_done = True
        
        # Payment 1 in words
        elif 'Viết bằng chữ' in text and payment_1_done and not payment_1_words_done:
            _replace_paragraph_text(p, f'Viết bằng chữ: {number_to_vietnamese_words(data.deposit_amount)}')
            payment_1_words_done = True
        
        # Payment 2 (remaining)
        elif 'Lần 2' in text:
            _replace_paragraph_text(
                p,
                f'Lần 2 : {format_vnd(data.remaining_amount)} đ khi kết thúc tiệc.'
            )
        
        # Payment 2 in words
        elif 'Viết bằng chữ' in text and payment_1_words_done:
            _replace_paragraph_text(p, f'Viết bằng chữ: {number_to_vietnamese_words(data.remaining_amount)}')
            break  # Done with payments

    # ── 8. Service checklist (P37-P44) — mark included services ──
    # These are static text lines; we leave them as-is from the template.

    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

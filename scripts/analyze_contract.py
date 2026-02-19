"""Deep analysis of contract template - find exact placeholder positions with run-level detail."""
from docx import Document

doc = Document(r'd:\PROJECT\AM THUC GIAO TUYET\HDDV Giao Tuyet Template.docx')

out = []

# Key paragraphs that need data replacement
key_paragraphs = {
    'contract_number': None,  # "Số :" in table header
    'customer_name': None,    # "Ông/Bà:"
    'customer_address': None, # "Địa chỉ:"
    'customer_id': None,      # "CMND/CCCD :"
    'customer_phone': None,   # "Điện thoại:" (first one)
    'party_time': None,       # "Thời gian đãi tiệc"
    'party_end': None,        # "Thời gian kết thúc tiệc"
    'party_address': None,    # "Địa chỉ đãi tiệc"
    'table_count': None,      # "Số Bàn Tiệc"
    'menu_items': None,       # "Thực Đơn : bao gồm..."
    'unit_price': None,       # "Đơn Giá :"
    'total_amount': None,     # "Tổng Cộng Thanh Toán :"
    'total_words': None,      # "Viết Bằng chữ :"
    'payment_1': None,        # "Lần 1 :"
    'payment_1_words': None,  # "Viết bằng chữ:" after Lần 1
    'payment_2': None,        # "Lần 2 :"
    'payment_2_words': None,  # "Viết bằng chữ:" after Lần 2
}

out.append("=" * 80)
out.append("RUN-LEVEL DETAIL FOR KEY PARAGRAPHS")
out.append("=" * 80)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    
    # Check if this is a key paragraph
    is_key = False
    for kw in ['Số :', 'Ông/Bà', 'Địa chỉ', 'CMND', 'Điện thoại', 'đãi tiệc', 
               'kết thúc tiệc', 'Bàn Tiệc', 'Thực Đơn', 'Đơn Giá', 'Tổng Cộng', 
               'Viết Bằng', 'Viết bằng', 'Lần 1', 'Lần 2', 'Hôm nay']:
        if kw in text:
            is_key = True
            break
    
    if is_key:
        out.append(f"\nP{i}: '{text}'")
        out.append(f"  Style: {p.style.name}")
        out.append(f"  Runs ({len(p.runs)}):")
        for ri, run in enumerate(p.runs):
            rtext = run.text
            out.append(f"    R{ri}: '{rtext}' | bold={run.bold} font={run.font.name} size={run.font.size}")

# Check Table 0 header for contract number
out.append("\n" + "=" * 80)
out.append("TABLE 0 - HEADER (Contract Number)")
out.append("=" * 80)
table = doc.tables[0]
for ri, row in enumerate(table.rows):
    for ci, cell in enumerate(row.cells):
        for pi, p in enumerate(cell.paragraphs):
            if p.text.strip():
                out.append(f"\nT0.R{ri}.C{ci}.P{pi}: '{p.text.strip()}'")
                for rri, run in enumerate(p.runs):
                    out.append(f"    R{rri}: '{run.text}' | bold={run.bold} font={run.font.name}")

# Check paragraphs between menu and pricing for dish items
out.append("\n" + "=" * 80)
out.append("PARAGRAPHS P25-P35 (Menu items area)")
out.append("=" * 80)
for i in range(25, 36):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        out.append(f"P{i}: '{p.text}' | style={p.style.name}")

# Write output
with open(r'd:\PROJECT\AM THUC GIAO TUYET\scripts\contract_deep_analysis.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(out))

print("Done! Output saved to scripts/contract_deep_analysis.txt")
